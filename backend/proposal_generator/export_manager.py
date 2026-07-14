

import json
import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict

from backend.proposal_generator.exceptions import ExportError
from backend.proposal_generator.proposal_builder import ProposalDraft

logger = logging.getLogger("fundforge.proposal_generator.export_manager")






@dataclass
class ExportResult:
    
    format:      str
    content:     str
    filename:    str
    mime_type:   str
    byte_size:   int           = 0
    exported_at: datetime      = field(default_factory=lambda: datetime.utcnow())

    def as_bytes(self) -> bytes:
        
        if self.format == "docx":
            import base64
            return base64.b64decode(self.content)
        return self.content.encode("utf-8")

    def to_dict(self) -> Dict[str, Any]:
        return {
            "format":      self.format,
            "filename":    self.filename,
            "mime_type":   self.mime_type,
            "byte_size":   self.byte_size,
            "exported_at": self.exported_at.isoformat(),
        }






class ExportManager:
    

    
    def export(
        self,
        draft:   ProposalDraft,
        format:  str = "markdown",
    ) -> ExportResult:
        
        fmt = format.lower().strip()
        try:
            if fmt in ("markdown", "md"):
                return self._to_markdown(draft)
            if fmt in ("txt", "text", "plain"):
                return self._to_plain_text(draft)
            if fmt == "json":
                return self._to_json(draft)
            if fmt in ("html", "htm"):
                return self._to_html(draft)
            if fmt == "docx":
                return self._to_docx(draft)
            raise ExportError(
                f"Unsupported export format '{format}'. "
                f"Supported: markdown, txt, json, html, docx.",
                export_format=format,
            )
        except ExportError:
            raise
        except Exception as exc:
            raise ExportError(
                f"Export to '{format}' failed: {exc}",
                export_format=format,
            ) from exc

    def to_pdf_ready_html(self, draft: ProposalDraft) -> str:
        
        return self._build_html(draft, print_mode=True)

    
    def _to_markdown(self, draft: ProposalDraft) -> ExportResult:
        content  = draft.full_text_md
        filename = self._safe_filename(draft, "md")
        return ExportResult(
            format    = "markdown",
            content   = content,
            filename  = filename,
            mime_type = "text/markdown; charset=utf-8",
            byte_size = len(content.encode("utf-8")),
        )

    def _to_plain_text(self, draft: ProposalDraft) -> ExportResult:
        content  = draft.full_text_plain
        filename = self._safe_filename(draft, "txt")
        return ExportResult(
            format    = "txt",
            content   = content,
            filename  = filename,
            mime_type = "text/plain; charset=utf-8",
            byte_size = len(content.encode("utf-8")),
        )

    def _to_json(self, draft: ProposalDraft) -> ExportResult:
        data = {
            "proposal_id":   draft.proposal_id,
            "version":       draft.version,
            "company_name":  draft.company_name,
            "grant_name":    draft.grant_name,
            "grant_id":      draft.grant_id,
            "template_id":   draft.template_id,
            "tone":          draft.tone,
            "generated_at":  draft.generated_at.isoformat(),
            "word_count":    draft.word_count,
            "completeness":  draft.completeness_pct,
            "sections":      {},
        }
        for key in draft.section_order:
            data["sections"][key] = draft.sections.get(key, "")

        content  = json.dumps(data, ensure_ascii=False, indent=2)
        filename = self._safe_filename(draft, "json")
        return ExportResult(
            format    = "json",
            content   = content,
            filename  = filename,
            mime_type = "application/json; charset=utf-8",
            byte_size = len(content.encode("utf-8")),
        )

    def _to_html(self, draft: ProposalDraft) -> ExportResult:
        content  = self._build_html(draft, print_mode=False)
        filename = self._safe_filename(draft, "html")
        return ExportResult(
            format    = "html",
            content   = content,
            filename  = filename,
            mime_type = "text/html; charset=utf-8",
            byte_size = len(content.encode("utf-8")),
        )

    def _to_docx(self, draft: ProposalDraft) -> ExportResult:
        import io
        import base64
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError as exc:
            raise ExportError(
                "python-docx is not installed. DOCX generation will be unavailable.",
                export_format="docx"
            ) from exc

        doc = DocxDocument()

        
        title_text = f"Grant Proposal: {draft.grant_name}"
        title = doc.add_heading(title_text, level=0)
        title.alignment = 0  

        
        meta_p = doc.add_paragraph()
        meta_p.add_run("Company Name: ").bold = True
        meta_p.add_run(f"{draft.company_name}   |   ")
        meta_p.add_run("Industry: ").bold = True
        meta_p.add_run(f"{draft.industry}   |   ")
        meta_p.add_run("Stage: ").bold = True
        meta_p.add_run(f"{draft.stage}   |   ")
        meta_p.add_run("Location: ").bold = True
        meta_p.add_run(f"{draft.location}\n")
        meta_p.add_run("Website: ").bold = True
        meta_p.add_run(f"{draft.website}   |   ")
        meta_p.add_run("Founding Year: ").bold = True
        meta_p.add_run(f"{draft.founding_year}   |   ")
        meta_p.add_run("Team Size: ").bold = True
        meta_p.add_run(f"{draft.team_size}   |   ")
        meta_p.add_run("Grant Name: ").bold = True
        meta_p.add_run(f"{draft.grant_name}\n")
        meta_p.add_run("Proposal Version: ").bold = True
        meta_p.add_run(f"{draft.version}   |   ")
        meta_p.add_run("Generation Date: ").bold = True
        meta_p.add_run(f"{draft.generated_at.strftime('%d %b %Y')}\n")
        meta_p.paragraph_format.space_after = Pt(18)

        
        for key in draft.section_order:
            content = draft.sections.get(key, "")
            if not content:
                continue
            
            section_title = key.replace("_", " ").upper()
            h = doc.add_heading(section_title, level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True

            
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(10)
                    p.paragraph_format.line_spacing = 1.15
        
        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        content_b64 = base64.b64encode(docx_bytes).decode("utf-8")
        filename = self._safe_filename(draft, "docx")
        
        return ExportResult(
            format    = "docx",
            content   = content_b64,
            filename  = filename,
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            byte_size = len(docx_bytes),
        )

    def _build_html(self, draft: ProposalDraft, print_mode: bool = False) -> str:
        
        pass

        
        md = draft.full_text_md
        html_body = self._md_to_html(md)

        page_size = "A4" if print_mode else "letter"
        margin    = "2.5cm" if print_mode else "2rem"

        css = f"""
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            font-size: 14px;
            line-height: 1.7;
            color: #1f2328;
            max-width: 820px;
            margin: 0 auto;
            padding: {margin};
        }}
        h1 {{ font-size: 22px; border-bottom: 2px solid #3b82d4; padding-bottom: 6px; }}
        h2 {{ font-size: 17px; color: #1d4ed8; margin-top: 28px; }}
        p  {{ margin: 10px 0; text-align: justify; }}
        hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 20px 0; }}
        .meta {{ color: #57606a; font-size: 13px; margin-bottom: 18px; }}
        @page {{ size: {page_size}; margin: 2cm; }}
        """

        company_safe = self._escape_html(draft.company_name)
        industry_safe = self._escape_html(draft.industry)
        stage_safe = self._escape_html(draft.stage)
        location_safe = self._escape_html(draft.location)
        website_safe = self._escape_html(draft.website)
        founding_year_safe = self._escape_html(draft.founding_year)
        team_size_safe = self._escape_html(draft.team_size)
        grant_safe   = self._escape_html(draft.grant_name)

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Grant Proposal — {company_safe}</title>
<style>{css}</style>
</head>
<body>
<div class="meta" style="margin-bottom: 25px; padding-bottom: 15px; border-bottom: 1px solid #e5e7eb; font-size: 13px; line-height: 1.6; color: #57606a;">
  <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 8px;">
    <div><strong>Company Name:</strong> {company_safe}</div>
    <div><strong>Industry:</strong> {industry_safe}</div>
    <div><strong>Stage:</strong> {stage_safe}</div>
    <div><strong>Location:</strong> {location_safe}</div>
    <div><strong>Website:</strong> {website_safe}</div>
    <div><strong>Founding Year:</strong> {founding_year_safe}</div>
    <div><strong>Team Size:</strong> {team_size_safe}</div>
    <div><strong>Grant Name:</strong> {grant_safe}</div>
    <div><strong>Proposal Version:</strong> {draft.version}</div>
    <div><strong>Generation Date:</strong> {draft.generated_at.strftime('%d %b %Y')}</div>
  </div>
</div>
{html_body}
</body>
</html>"""

    
    @staticmethod
    def _md_to_html(text: str) -> str:
        
        import re

        lines   = text.split("\n")
        result  = []
        buffer  = []

        def flush():
            if buffer:
                para = " ".join(buffer).strip()
                if para:
                    result.append(f"<p>{para}</p>")
                buffer.clear()

        for line in lines:
            stripped = line.strip()

            if not stripped:
                flush()
                continue

            if stripped.startswith("# "):
                flush()
                result.append(f"<h1>{ExportManager._escape_html(stripped[2:])}</h1>")
            elif stripped.startswith("## "):
                flush()
                result.append(f"<h2>{ExportManager._escape_html(stripped[3:])}</h2>")
            elif stripped.startswith("### "):
                flush()
                result.append(f"<h3>{ExportManager._escape_html(stripped[4:])}</h3>")
            elif stripped == "---":
                flush()
                result.append("<hr>")
            else:
                inline = ExportManager._escape_html(stripped)
                inline = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", inline)
                inline = re.sub(r"\*(.+?)\*",    r"<em>\1</em>", inline)
                buffer.append(inline)

        flush()
        return "\n".join(result)

    @staticmethod
    def _escape_html(text: str) -> str:
        
        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    @staticmethod
    def _safe_filename(draft: ProposalDraft, ext: str) -> str:
        
        import re
        company = re.sub(r"[^\w\s-]", "", draft.company_name).strip()
        company = re.sub(r"[\s]+", "_", company)[:30]
        grant   = re.sub(r"[^\w\s-]", "", draft.grant_id or "grant").strip()
        grant   = re.sub(r"[\s]+", "_", grant)[:20]
        return f"{company}_{grant}_v{draft.version}.{ext}"
